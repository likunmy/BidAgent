import io
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from PIL import Image, ImageDraw, ImageFont
from pydantic import BaseModel
from sqlalchemy.orm import Session

from app.agent.tools import image_extractor
from app.agent.tools.missing_info import MissingItem, store_missing_info, StoreMissingInfoInput
from app.core.config import settings


class ParagraphBlock(BaseModel):
    type: str = "paragraph"
    text: str


class ImageBlock(BaseModel):
    type: str = "image"
    source_md_path: str
    md5: str
    width: int
    height: int


class TableBlock(BaseModel):
    type: str = "table"
    headers: list[str]
    rows: list[list[str]]


class PlaceholderBlock(BaseModel):
    type: str = "placeholder"
    name: str
    description: str | None = None
    width: int = 400
    height: int = 250


class Section(BaseModel):
    heading: str
    level: int = 1
    children: list[ParagraphBlock | ImageBlock | TableBlock | PlaceholderBlock] = []


class OutputDocxInput(BaseModel):
    project_id: int
    output_filename: str
    sections: list[Section]


class OutputDocxOutput(BaseModel):
    file_path: str
    missing_items: list[MissingItem]


def output_docx(input_data: OutputDocxInput, db: Session) -> OutputDocxOutput:
    """Generate a docx file from structured section content.

    Supports paragraph, image (with AI-controlled size), table, and
    placeholder blocks. Placeholders generate a placeholder image in
    the docx and store a MissingInfo record.
    """
    doc = Document()
    missing_items: list[MissingItem] = []

    for section in input_data.sections:
        _add_heading(doc, section.heading, section.level)

        for child in section.children:
            if child.type == "paragraph":
                _add_paragraph(doc, child.text)

            elif child.type == "image":
                _add_image(doc, child)

            elif child.type == "table":
                _add_table(doc, child)

            elif child.type == "placeholder":
                _add_placeholder(doc, child)
                missing_items.append(MissingItem(
                    name=child.name,
                    description=child.description,
                ))

    # Save
    output_dir = Path(settings.upload_dir) / "projects" / str(input_data.project_id) / "output"
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / input_data.output_filename
    doc.save(str(output_path))

    # Store missing info
    if missing_items:
        store_missing_info(StoreMissingInfoInput(
            project_id=input_data.project_id,
            items=missing_items,
        ), db)

    return OutputDocxOutput(
        file_path=str(output_path),
        missing_items=missing_items,
    )


def _add_heading(doc: Document, text: str, level: int):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)


def _add_paragraph(doc: Document, text: str):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)


def _add_image(doc: Document, block: ImageBlock):
    img_bytes = image_extractor.get_image_bytes(block.source_md_path, block.md5)
    if img_bytes:
        stream = io.BytesIO(img_bytes)
        doc.add_picture(stream, width=Inches(block.width / 96), height=Inches(block.height / 96))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_table(doc: Document, block: TableBlock):
    rows = 1 + len(block.rows)
    cols = max(len(block.headers), *(len(r) for r in block.rows)) if block.rows else len(block.headers)
    table = doc.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(block.headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)

    # Data rows
    for ri, row_data in enumerate(block.rows):
        for ci, val in enumerate(row_data):
            if ci < cols:
                table.rows[ri + 1].cells[ci].text = val


def _add_placeholder(doc: Document, block: PlaceholderBlock):
    """Insert a placeholder image into the docx."""
    img_bytes = _generate_placeholder_image(block.name, block.width, block.height)
    stream = io.BytesIO(img_bytes)
    doc.add_picture(stream, width=Inches(block.width / 96), height=Inches(block.height / 96))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add caption below
    cap = doc.add_paragraph(f"[缺失: {block.name}]")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)


def _generate_placeholder_image(name: str, width: int, height: int) -> bytes:
    """Generate a placeholder image with the missing item name overlaid."""
    img = Image.new("RGB", (width, height), color=(245, 240, 235))
    draw = ImageDraw.Draw(img)

    # Try to load a CJK-friendly font, fall back to default
    try:
        font = ImageFont.truetype("simhei.ttf", max(12, min(width, height) // 12))
    except Exception:
        try:
            font = ImageFont.truetype("msyh.ttc", max(12, min(width, height) // 12))
        except Exception:
            font = ImageFont.load_default()

    # Draw border
    draw.rectangle([(0, 0), (width - 1, height - 1)], outline=(200, 195, 185), width=2)

    # Centered text
    text = f"缺失: {name}"
    bbox = draw.textbbox((0, 0), text, font=font)
    text_w = bbox[2] - bbox[0]
    text_h = bbox[3] - bbox[1]
    x = (width - text_w) / 2
    y = (height - text_h) / 2
    draw.text((x, y), text, fill=(160, 150, 135), font=font)

    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()
